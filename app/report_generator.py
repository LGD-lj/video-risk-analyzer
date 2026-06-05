"""Word 报告生成 —— 使用 python-docx"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from .models import VideoInfo, RiskPoint


def _format_duration(seconds: float) -> str:
    """格式化时长"""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    if h > 0:
        return f"{h}时{m}分{s}秒"
    return f"{m}分{s}秒"


def _severity_color(severity: str) -> RGBColor:
    """风险等级对应颜色"""
    if severity == "高":
        return RGBColor(0xCC, 0x00, 0x00)  # 红色
    elif severity == "中":
        return RGBColor(0xE6, 0x8A, 0x00)  # 橙色
    else:
        return RGBColor(0x00, 0x80, 0x00)  # 绿色


def generate_word_report(
    video_info: VideoInfo,
    risk_points: list[RiskPoint],
    output_path: str,
) -> str:
    """生成 Word 报告"""
    doc = Document()

    # ---------- 页面设置 ----------
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ==================== 封面/标题 ====================
    # 空行
    for _ in range(3):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("视频风险点分析报告")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("行车记录视频 — AI 智能风险识别")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # 封面的基本信息
    info_cover = doc.add_paragraph()
    info_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_cover.add_run(f"视频文件：{video_info.filename}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    info_cover2 = doc.add_paragraph()
    info_cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_cover2.add_run(
        f"时长：{_format_duration(video_info.duration_seconds)}　|　"
        f"分辨率：{video_info.resolution}　|　"
        f"帧率：{video_info.fps:.1f} fps"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 分页
    doc.add_page_break()

    # ==================== 一、视频基本信息 ====================
    doc.add_heading("一、视频基本信息", level=1)

    info_table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ("文件名", video_info.filename),
        ("视频时长", _format_duration(video_info.duration_seconds)),
        ("分辨率", video_info.resolution),
        ("帧率 (fps)", f"{video_info.fps:.2f}"),
        ("编码格式", video_info.codec),
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        # 第一列加粗
        for paragraph in info_table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.add_paragraph("")

    # ==================== 二、风险点汇总 ====================
    doc.add_heading("二、风险点汇总", level=1)

    summary_text = f"本次分析共发现 {len(risk_points)} 个风险点，详情如下："
    p = doc.add_paragraph(summary_text)
    p.runs[0].font.size = Pt(11)

    # 汇总表
    summary_table = doc.add_table(rows=len(risk_points) + 1, cols=4, style="Light Grid Accent 1")
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ["序号", "时间点", "风险等级", "风险类型"]
    for j, header in enumerate(headers):
        cell = summary_table.rows[0].cells[j]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # 表内容
    for i, rp in enumerate(risk_points):
        row = summary_table.rows[i + 1]
        row.cells[0].text = str(i + 1)
        row.cells[1].text = rp.timestamp_display
        row.cells[2].text = rp.severity.value
        row.cells[3].text = "、".join([rt.value for rt in rp.risk_types])

        # 风险等级着色
        color = _severity_color(rp.severity.value)
        for paragraph in row.cells[2].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = color
                run.font.bold = True

    doc.add_paragraph("")

    # 统计
    high_count = sum(1 for rp in risk_points if rp.severity.value == "高")
    medium_count = sum(1 for rp in risk_points if rp.severity.value == "中")
    low_count = sum(1 for rp in risk_points if rp.severity.value == "低")

    stats = doc.add_paragraph()
    run = stats.add_run(
        f"风险统计：高风险 {high_count} 个　|　"
        f"中风险 {medium_count} 个　|　"
        f"低风险 {low_count} 个"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 分页
    doc.add_page_break()

    # ==================== 三、风险点详情 ====================
    doc.add_heading("三、风险点详情", level=1)

    for i, rp in enumerate(risk_points):
        # 风险点标题
        heading_text = f"风险点 {i + 1}：{rp.timestamp_display} — {rp.severity.value}风险"
        doc.add_heading(heading_text, level=2)

        # 详情表
        detail_table = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
        detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        detail_data = [
            ("时间点", rp.timestamp_display),
            ("风险等级", rp.severity.value),
            ("风险类型", "、".join([rt.value for rt in rp.risk_types])),
            ("风险描述", rp.description),
        ]

        for j, (label, value) in enumerate(detail_data):
            detail_table.rows[j].cells[0].text = label
            detail_table.rows[j].cells[1].text = value
            for paragraph in detail_table.rows[j].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

            # 风险等级着色
            if label == "风险等级":
                color = _severity_color(value)
                for paragraph in detail_table.rows[j].cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = color
                        run.font.bold = True

        # 插入截图（如果存在）
        screenshot_path = rp.screenshot_path
        if os.path.exists(screenshot_path):
            doc.add_paragraph("")  # 空行
            img_paragraph = doc.add_paragraph()
            img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                run = img_paragraph.add_run()
                run.add_picture(screenshot_path, width=Inches(5.0))
                # 图片说明
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(f"图 {i+1}：{rp.timestamp_display} 画面截图")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                run.font.italic = True
            except Exception:
                doc.add_paragraph("（截图插入失败）")

        # 风险点之间加分隔（最后一个不加）
        if i < len(risk_points) - 1:
            doc.add_paragraph("─" * 50)

    # 分页
    doc.add_page_break()

    # ==================== 四、备注 ====================
    doc.add_heading("四、备注与声明", level=1)

    notes = [
        "1. 本报告由 AI 视觉模型自动分析生成，分析结果可能存在偏差，仅供参考。",
        "2. 视频清晰度、天气、光照等因素可能影响识别准确性。",
        "3. 对于限高标识，AI 可能无法准确识别具体的限高数值，请在实际行驶中以现场标识为准，必要时进行人工复核。",
        "4. 低净空风险判定基于画面中桥梁/隧道的高度视觉估计，不适用于所有车型。",
        "5. 建议结合本报告与实际路况进行综合判断，确保行车安全。",
        "6. 本系统生成的截图经过 JPEG 压缩，可能损失部分细节。",
    ]

    for note in notes:
        p = doc.add_paragraph(note)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ---------- 保存 ----------
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
